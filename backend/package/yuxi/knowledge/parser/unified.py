"""Unified parser module for markdown conversion."""

from __future__ import annotations

import asyncio
import base64
import os
import re
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import aiofiles
from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter
from langchain_community.document_loaders import PyPDFLoader
from markdownify import markdownify as md_convert

from yuxi.knowledge.parser.zip_utils import process_zip_file as _process_zip_file
from yuxi.storage.minio import get_minio_client
from yuxi.utils import logger

AUDIO_EXTENSIONS = (".mp3", ".wav", ".m4a", ".flac", ".aac", ".ogg", ".wma")
VIDEO_EXTENSIONS = (".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm")
IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif")

SUPPORTED_FILE_EXTENSIONS: tuple[str, ...] = (
    ".txt",
    ".md",
    ".docx",
    ".html",
    ".htm",
    ".json",
    ".csv",
    ".xls",
    ".xlsx",
    ".pdf",
    ".pptx",
    *IMAGE_EXTENSIONS,
    ".zip",
    *AUDIO_EXTENSIONS,
    *VIDEO_EXTENSIONS,
)


def is_supported_file_extension(file_name: str | os.PathLike[str]) -> bool:
    """Check whether the given file path has a supported extension."""
    return Path(file_name).suffix.lower() in SUPPORTED_FILE_EXTENSIONS


@dataclass(slots=True)
class MarkdownParseResult:
    """统一的 Markdown 解析结果。"""

    markdown: str
    file_ext: str | None = None
    artifacts: dict[str, Any] = field(default_factory=dict)


_docling_converter: DocumentConverter | None = None


def _get_docling_converter() -> DocumentConverter:
    """获取 Docling 文档转换器单例。"""
    global _docling_converter
    if _docling_converter is None:
        _docling_converter = DocumentConverter(
            format_options={
                InputFormat.DOCX: None,
                InputFormat.XLSX: None,
                InputFormat.PPTX: None,
            }
        )
    return _docling_converter


def _resolve_image_storage_params(params: dict | None) -> tuple[str, str]:
    params = params or {}

    image_bucket = params.get("image_bucket") or "public"
    image_prefix = params.get("image_prefix")
    if image_prefix:
        normalized_prefix = str(image_prefix).strip("/")
        if normalized_prefix:
            return image_bucket, normalized_prefix

    return image_bucket, "unknown/kb-images"


def _resolve_ocr_engine_params(params: dict | None) -> tuple[str, dict[str, Any]]:
    from yuxi import config

    params = params or {}
    engine = str(params.get("ocr_engine") if "ocr_engine" in params else config.default_ocr_engine)
    engine = engine.strip() or config.default_ocr_engine
    engine_config = params.get("ocr_engine_config")
    processor_params = dict(params)
    if isinstance(engine_config, dict):
        processor_params.update(engine_config)
    return engine, processor_params


def _upload_image_to_minio(image_data: bytes, filename: str, bucket_name: str, object_prefix: str) -> str:
    """上传图片到 MinIO，返回 URL。"""
    minio_client = get_minio_client()
    minio_client.ensure_bucket_exists(bucket_name)

    normalized_prefix = object_prefix.strip("/") or "unknown/kb-images"
    timestamp = int(time.time() * 1000000)
    object_name = f"{normalized_prefix}/{timestamp}_{Path(filename).name}"

    result = minio_client.upload_file(
        bucket_name=bucket_name,
        object_name=object_name,
        data=image_data,
    )
    return result.url


def _parse_data_uri(data_uri: str) -> tuple[bytes, str]:
    """解析 data URI，返回 (image_data, mime_type)。"""
    header, base64_data = data_uri.split(",", 1)
    mime_type = header.split(":")[1].split(";")[0]
    image_data = base64.b64decode(base64_data)
    return image_data, mime_type


def _convert_with_docling(file_path: Path, params: dict | None = None) -> str:
    """使用 Docling 将 docx/xlsx/pptx 转换为 Markdown。"""
    params = params or {}
    image_bucket, image_prefix = _resolve_image_storage_params(params)

    converter = _get_docling_converter()
    result = converter.convert(file_path)

    if result.status.name != "SUCCESS":
        raise RuntimeError(f"Docling 转换失败: {result.status}")

    doc = result.document

    if hasattr(doc, "pictures") and doc.pictures:
        replacements: list[str] = []
        for pic in doc.pictures:
            uri = str(pic.image.uri) if hasattr(pic, "image") and hasattr(pic.image, "uri") else ""
            if uri.startswith("data:"):
                filename = "image"
                try:
                    image_data, mime_type = _parse_data_uri(uri)
                    filename = f"image_{int(time.time() * 1000000)}.{mime_type.split('/')[-1]}"
                    url = _upload_image_to_minio(image_data, filename, image_bucket, image_prefix)
                    replacements.append(f"![{filename}]({url})")
                except Exception as e:  # noqa: BLE001
                    logger.error(f"上传图片失败 {filename}: {e}")
                    replacements.append(f"[图片: {filename}]")
            else:
                replacements.append("")

        markdown = doc.export_to_markdown()
        for replacement in replacements:
            markdown = re.sub(r"<!--\s*image\s*-->", replacement, markdown, count=1)
        return markdown

    return doc.export_to_markdown()


def _read_text_best_effort(file_path: Path) -> str:
    """健壮读取纯文本：按常见编码依次尝试，兜底 latin-1 replace。

    修复：原先硬编码 utf-8，GBK/GB18030/Big5 等非 UTF-8 中文文本会 UnicodeDecodeError。
    """
    raw = file_path.read_bytes()
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")
    for enc in ("utf-8", "gb18030", "gbk", "big5", "utf-16"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    try:
        from charset_normalizer import from_bytes

        best = from_bytes(raw).best()
        if best is not None:
            return str(best)
    except Exception:  # noqa: BLE001
        pass
    return raw.decode("latin-1", errors="replace")


def _convert_docx_with_python_docx(file_path: Path) -> str:
    """使用 python-docx 解析 DOCX（Docling 失败时兜底）。"""
    from docx import Document

    document = Document(str(file_path))
    blocks: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        header = rows[0]
        blocks.append(f"| {' | '.join(header)} |")
        blocks.append(f"| {' | '.join(['---'] * len(header))} |")

        for row in rows[1:]:
            normalized_row = row + [""] * (len(header) - len(row))
            blocks.append(f"| {' | '.join(normalized_row[: len(header)])} |")

        blocks.append("")

    return "\n\n".join(blocks).strip()


def pdfreader(file_path, params=None):
    """读取 PDF 文件并返回 text 文本。"""
    if isinstance(file_path, str):
        file_path = Path(file_path)

    assert file_path.exists(), "File not found"
    assert file_path.suffix.lower() == ".pdf", "File format not supported"

    loader = PyPDFLoader(str(file_path))
    docs = loader.load()
    text = "\n\n".join([d.page_content for d in docs])
    return text


def parse_pdf(file, params=None):
    """解析 PDF 文件，支持多种 OCR 方式。"""
    from yuxi.knowledge.parser.base import DocumentProcessorException
    from yuxi.knowledge.parser.factory import DocumentProcessorFactory

    opt_ocr, processor_params = _resolve_ocr_engine_params(params)

    if opt_ocr == "disable":
        return pdfreader(file, params=processor_params)

    image_bucket, image_prefix = _resolve_image_storage_params(processor_params)
    processor_params.setdefault("image_bucket", image_bucket)
    processor_params.setdefault("image_prefix", image_prefix)

    try:
        if processor_params.get("ocr_fallback"):
            return DocumentProcessorFactory.process_file_with_fallback(
                opt_ocr, file, processor_params, processor_params.get("ocr_fallback_chain")
            )
        return DocumentProcessorFactory.process_file(opt_ocr, file, processor_params)
    except DocumentProcessorException as e:
        logger.error(f"文档处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:  # noqa: BLE001
        logger.error(f"PDF 解析失败: {str(e)}")
        raise DocumentProcessorException(f"PDF解析失败: {str(e)}", opt_ocr, "parsing_failed")


def parse_image(file, params=None):
    """解析图像文件，支持多种 OCR 方式。"""
    from yuxi.knowledge.parser.base import DocumentProcessorException
    from yuxi.knowledge.parser.factory import DocumentProcessorFactory

    opt_ocr, processor_params = _resolve_ocr_engine_params(params)

    if opt_ocr == "disable":
        raise ValueError(
            "图像文件必须启用OCR才能提取文本内容。"
            "请选择OCR方式 "
            "(rapid_ocr/mineru_ocr/mineru_official/pp_structure_v3_ocr/deepseek_ocr/"
            "paddleocr_vl_1_6/paddleocr_pp_ocrv6) 或移除该文件。"
        )

    image_bucket, image_prefix = _resolve_image_storage_params(processor_params)
    processor_params.setdefault("image_bucket", image_bucket)
    processor_params.setdefault("image_prefix", image_prefix)

    try:
        if processor_params.get("ocr_fallback"):
            return DocumentProcessorFactory.process_file_with_fallback(
                opt_ocr, file, processor_params, processor_params.get("ocr_fallback_chain")
            )
        return DocumentProcessorFactory.process_file(opt_ocr, file, processor_params)
    except DocumentProcessorException as e:
        logger.error(f"图像处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:  # noqa: BLE001
        logger.error(f"图像解析失败: {str(e)}")
        raise DocumentProcessorException(f"图像解析失败: {str(e)}", opt_ocr, "parsing_failed")


async def parse_pdf_async(file, params=None):
    return await asyncio.to_thread(parse_pdf, file, params=params)


async def parse_image_async(file, params=None):
    return await asyncio.to_thread(parse_image, file, params=params)


async def _parse_image_with_vl(file_path: str, params: dict | None = None) -> str:
    """
    图片解析：OCR 文字提取（可选）+ VL 画面描述（默认开启），合并为 markdown

    无文字图片靠 VL 描述保证可检索；结果写入 params["_media_preprocess"]
    供打标链路缓存复用，避免重复调用多模态模型。
    """
    params = params if params is not None else {}
    opt_ocr, _ = _resolve_ocr_engine_params(params)
    enable_vl = params.get("enable_image_vl", True)

    ocr_text = ""
    ocr_error: Exception | None = None
    if opt_ocr != "disable":
        try:
            ocr_text = (await parse_image_async(file_path, params=params)) or ""
        except Exception as e:  # noqa: BLE001
            ocr_error = e
            logger.warning(f"图片 OCR 失败，尝试 VL 描述兜底: {e}")

    vl_text = ""
    vl_model = ""
    if enable_vl:
        try:
            from yuxi.services.tagging.preprocessors import ImagePreprocessor
            from yuxi.services.tagging.prompt_config import get_flat_processing_config

            async with aiofiles.open(file_path, "rb") as f:
                image_data = await f.read()
            processor = ImagePreprocessor(get_flat_processing_config())
            vl_result = await processor.process(image_data, Path(file_path).name)
            vl_text = (vl_result.text or "").strip()
            vl_model = vl_result.model_used
        except Exception as e:  # noqa: BLE001
            logger.warning(f"图片 VL 描述失败: {e}")

    ocr_text = ocr_text.strip()
    if not ocr_text and not vl_text:
        if ocr_error is not None:
            raise ocr_error
        raise ValueError(
            "图片解析失败：OCR 与 VL 描述均未产出内容。"
            "请启用 OCR (ocr_engine) 或检查 VL 模型配置 (enable_image_vl)。"
        )

    parts = []
    if ocr_text:
        parts.append(f"## 图片文字（OCR）\n\n{ocr_text}")
    if vl_text:
        parts.append(f"## 画面描述（VL）\n\n{vl_text}")
    combined = "\n\n".join(parts)

    params["_media_preprocess"] = {
        "text": combined[:10000],
        "source_type": "image",
        "model_used": vl_model or "ocr-only",
        "metadata": {"has_ocr": bool(ocr_text), "has_vl": bool(vl_text)},
    }
    return combined


async def _process_file_to_markdown_core(
    file_path: str, params: dict | None = None
) -> tuple[str, str | None, dict[str, Any]]:
    """将不同类型的文件转换为 markdown，支持本地文件和 MinIO 文件。"""
    from yuxi.knowledge.utils.kb_utils import is_minio_url, parse_minio_url
    from yuxi.storage.minio.client import get_minio_client

    if is_minio_url(file_path):
        logger.debug(f"Downloading file from MinIO: {file_path}")

        if "?" in file_path:
            file_path_clean = file_path.split("?")[0]
        else:
            file_path_clean = file_path

        original_filename = file_path_clean.split("/")[-1]

        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(original_filename).suffix) as temp_file:
            temp_path = temp_file.name

        try:
            bucket_name, object_name = parse_minio_url(file_path)
            minio_client = get_minio_client()
            file_content = await minio_client.adownload_file(bucket_name, object_name)

            async with aiofiles.open(temp_path, "wb") as f:
                await f.write(file_content)

            logger.debug(f"File downloaded to temp path: {temp_path}")
            actual_file_path = temp_path

        except Exception as e:  # noqa: BLE001
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            logger.error(f"Failed to download file from MinIO: {e}")
            raise ValueError(f"无法从MinIO下载文件: {e}")
    else:
        actual_file_path = file_path

    file_ext: str | None = None
    artifacts: dict[str, Any] = {}

    try:
        file_path_obj = Path(actual_file_path)
        file_ext = file_path_obj.suffix.lower()

        if file_ext == ".pdf":
            text = await parse_pdf_async(str(file_path_obj), params=params)
            result = f"{text}"

        elif file_ext in [".txt", ".md"]:
            content = _read_text_best_effort(file_path_obj)
            result = f"{content}"

        elif file_ext == ".docx":
            try:
                result = _convert_with_docling(file_path_obj, params=params)
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Docling 解析 DOCX 失败，回退到 python-docx: {file_path_obj.name}, {e}")
                result = _convert_docx_with_python_docx(file_path_obj)

        elif file_ext == ".pptx":
            result = _convert_with_docling(file_path_obj, params=params)

        elif file_ext == ".doc":
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader

            loader = UnstructuredWordDocumentLoader(str(file_path_obj))
            docs = loader.load()
            result = "\n".join(doc.page_content for doc in docs).strip()

        elif file_ext in IMAGE_EXTENSIONS:
            # OCR 文字提取 + VL 画面描述合并（无文字图片靠 VL 保证可检索）
            result = await _parse_image_with_vl(str(file_path_obj), params=params)

        elif file_ext in [".html", ".htm"]:
            content = _read_text_best_effort(file_path_obj)
            text = md_convert(content, heading_style="ATX")
            result = f"{text}"

        elif file_ext == ".csv":
            import pandas as pd

            df = pd.read_csv(file_path_obj)
            markdown_content = ""

            for _, row in df.iterrows():
                row_df = pd.DataFrame([row], columns=df.columns)
                markdown_table = row_df.to_markdown(index=False)
                markdown_content += f"{markdown_table}\n\n"

            result = markdown_content.strip()

        elif file_ext in [".xls", ".xlsx"]:
            result = _convert_with_docling(file_path_obj, params=params)

        elif file_ext == ".json":
            import json

            async with aiofiles.open(file_path_obj, encoding="utf-8") as f:
                content = await f.read()
            data = json.loads(content)
            json_str = json.dumps(data, ensure_ascii=False, indent=2)
            result = f"```json\n{json_str}\n```"

        elif file_ext == ".zip":
            image_bucket, image_prefix = _resolve_image_storage_params(params)
            zip_result = await _process_zip_file(
                str(file_path_obj),
                image_bucket=image_bucket,
                image_prefix=image_prefix,
            )

            artifacts = {
                "zip_images_info": zip_result["images_info"],
                "zip_content_hash": zip_result["content_hash"],
                "zip_image_bucket": image_bucket,
                "zip_image_prefix": image_prefix,
            }

            result = zip_result["markdown_content"]

        elif file_ext in AUDIO_EXTENSIONS:
            # 音频文件：使用打标预处理器提取文本（路径直传，避免全量读内存）
            from yuxi.knowledge.utils.kb_utils import is_minio_url as _is_minio
            from yuxi.services.tagging.preprocessors import AudioPreprocessor
            from yuxi.services.tagging.prompt_config import get_flat_processing_config

            config = get_flat_processing_config()
            processor = AudioPreprocessor(config)
            original_filename = Path(file_path.split("?")[0]).name
            # 原始 MinIO URL 公网可达时，ASR 可走 filetrans 异步文件转写
            source_url = file_path if _is_minio(file_path) else None
            preprocess_result = await processor.process_path(
                str(file_path_obj), original_filename, source_url=source_url
            )
            result = preprocess_result.text

            if params is not None:
                # text 字段必须存在，打标链路缓存校验依赖它（缺失 = 重复预处理双倍 API 费用）
                params["_media_preprocess"] = {
                    "text": preprocess_result.text[:10000],
                    "source_type": preprocess_result.source_type,
                    "model_used": preprocess_result.model_used,
                    "metadata": preprocess_result.metadata,
                }

        elif file_ext in VIDEO_EXTENSIONS:
            # 视频文件：提取音轨+抽帧，合并文本
            from yuxi.services.tagging.preprocessors import VideoPreprocessor
            from yuxi.services.tagging.prompt_config import get_flat_processing_config

            config = get_flat_processing_config()
            processor = VideoPreprocessor(config)
            original_filename = Path(file_path.split("?")[0]).name
            # 注：视频音轨是 ffmpeg 重新分离出的本地 wav，无对应公网 URL，不传 source_url
            preprocess_result = await processor.process_path(str(file_path_obj), original_filename)
            result = preprocess_result.text

            if params is not None:
                params["_media_preprocess"] = {
                    "text": preprocess_result.text[:10000],
                    "source_type": preprocess_result.source_type,
                    "model_used": preprocess_result.model_used,
                    "metadata": preprocess_result.metadata,
                }

        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    except Exception:
        if is_minio_url(file_path) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as cleanup_e:  # noqa: BLE001
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {cleanup_e}")
        raise

    finally:
        if is_minio_url(file_path) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {e}")

    return result, file_ext, artifacts


async def parse_source_to_markdown(source: str, params: dict | None = None) -> MarkdownParseResult:
    """统一入口: 将文件解析为 Markdown（URL 解析已废弃）。"""
    markdown, file_ext, artifacts = await _process_file_to_markdown_core(source, params=params)
    return MarkdownParseResult(
        markdown=markdown,
        file_ext=file_ext,
        artifacts=artifacts,
    )


class Parser:
    """Lightweight facade for converting file sources to markdown."""

    @staticmethod
    async def aparse(source: str, params: dict | None = None) -> str:
        """Asynchronously parse source content and return markdown text."""
        parsed = await parse_source_to_markdown(source=source, params=params)
        return parsed.markdown

    @classmethod
    def parse(cls, source: str, params: dict | None = None) -> str:
        """Synchronously parse source content and return markdown text."""
        try:
            asyncio.get_running_loop()
        except RuntimeError:
            return asyncio.run(cls.aparse(source=source, params=params))

        raise RuntimeError("当前处于异步上下文，请使用 `await Parser.aparse(...)`")
